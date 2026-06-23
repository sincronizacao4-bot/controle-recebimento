"""
Geração do Controle de Recebimento em DOCX.
"""

import io
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

_BRASAO = os.path.join(os.path.dirname(os.path.abspath(__file__)), "brasao.png")

# Cores padrão
_AZUL       = "1A3F7A"   # cabeçalho das tabelas (azul escuro — tom da logo)
_LARANJA    = RGBColor(0xC0, 0x50, 0x00)   # título principal
_CINZA_HDR  = "D9D9D9"   # célula de label
_BRANCO     = RGBColor(0xFF, 0xFF, 0xFF)

# Largura total das tabelas (A4 – margens 2cm cada lado = 17cm)
_TBL_W  = Cm(17.0)
_COL_L  = Cm(5.0)   # coluna label (esquerda)
_COL_V  = Cm(12.0)  # coluna valor (direita)


# ── Helpers de célula/tabela ───────────────────────────────────────────────────
def _set_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell(cell, text: str, bold=False, size=9, center=False, white=False, italic=False):
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = _BRANCO
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)


def _header_row(row, cols: list[str]):
    for i, label in enumerate(cols):
        _set_bg(row.cells[i], _AZUL)
        _cell(row.cells[i], label, bold=True, center=True, white=True)


def _lock_table_width(tbl, total):
    """Trava a largura total da tabela para não esticar/encolher."""
    tbl_elem = tbl._tbl
    tblPr = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tblPr)
    # Remove tblW anterior se existir
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    # total é Cm — converte para twips (1 cm = 567 twips)
    twips = int(total.cm * 567)
    tblW.set(qn("w:w"), str(twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _info_table(doc, rows: list[tuple], header: str | None = None,
                col_widths=(_COL_L, _COL_V)):
    extra = 1 if header else 0
    tbl = doc.add_table(rows=len(rows) + extra, cols=2)
    tbl.style = "Table Grid"
    tbl.autofit = False
    _lock_table_width(tbl, _TBL_W)

    for row in tbl.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    if header:
        merged = tbl.rows[0].cells[0].merge(tbl.rows[0].cells[1])
        _set_bg(merged, _AZUL)
        _cell(merged, header, bold=True, white=True, center=True)

    for i, (label, value) in enumerate(rows):
        r = tbl.rows[i + extra]
        _set_bg(r.cells[0], _CINZA_HDR)
        _cell(r.cells[0], label, bold=True)
        _cell(r.cells[1], value)

    return tbl


def _spacer(doc, pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(pts)


# ── Cabeçalho ─────────────────────────────────────────────────────────────────
def _add_header(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    for p in header.paragraphs:
        p.clear()

    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)

    # Brasão
    if os.path.exists(_BRASAO):
        run_img = p.add_run()
        run_img.add_picture(_BRASAO, height=Cm(1.6))
        p.add_run("   ")

    # Bloco de texto ao lado
    run_nome = p.add_run("Prefeitura Municipal de Maracanaú\n")
    run_nome.bold = True
    run_nome.font.size = Pt(12)
    run_nome.font.color.rgb = RGBColor(0x1A, 0x3F, 0x7A)

    run_sec = p.add_run("Secretaria de Infraestrutura, Mobilidade e Controle Urbano\n")
    run_sec.font.size = Pt(9)
    run_sec.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    run_end = p.add_run("Avenida Durval Tomaz de Souza Nº 150 – Jereissati | Maracanaú")
    run_end.font.size = Pt(8)
    run_end.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Linha separadora
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1A3F7A")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Rodapé ─────────────────────────────────────────────────────────────────────
def _add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p.clear()

    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)

    # Linha separadora
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), "1A3F7A")
    pBdr.append(top)
    pPr.append(pBdr)

    run_addr = p.add_run(
        "Prefeitura Municipal de Maracanaú  |  "
        "Av. Durval Tomaz de Souza, Nº 150 – Jereissati  |  "
        "Maracanaú – CE  |  CEP 61.939-160"
    )
    run_addr.font.size = Pt(7)
    run_addr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Linha com número de página
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)

    def _fld(run, instrucao):
        fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
        it  = OxmlElement("w:instrText"); it.text = instrucao
        fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end")
        run._r.append(fc1); run._r.append(it); run._r.append(fc2)

    r_pg   = p2.add_run("Página "); r_pg.font.size = Pt(7); r_pg.font.color.rgb = RGBColor(0x66,0x66,0x66)
    r_num  = p2.add_run("");       r_num.font.size = Pt(7); r_num.font.color.rgb = RGBColor(0x66,0x66,0x66); _fld(r_num, "PAGE")
    r_de   = p2.add_run(" de ");   r_de.font.size  = Pt(7); r_de.font.color.rgb  = RGBColor(0x66,0x66,0x66)
    r_tot  = p2.add_run("");       r_tot.font.size = Pt(7); r_tot.font.color.rgb = RGBColor(0x66,0x66,0x66); _fld(r_tot, "NUMPAGES")


# ── Marca d'água ───────────────────────────────────────────────────────────────
def _add_watermark(doc):
    """
    Marca d'água centralizada no meio da página A4, repetida em todas as páginas.
    Usa VML posicionado absolutamente a partir do cabeçalho com offset calculado
    para cair no centro da folha (A4 = 841pt; desconta topo ~85pt → center ~336pt).
    """
    if not os.path.exists(_BRASAO):
        return
    try:
        section = doc.sections[0]
        hdr = section.header
        rId, _ = hdr.part.get_or_add_image(_BRASAO)

        # 340pt × 340pt, deslocado 336pt para baixo = centro visual da A4
        WMK_XML = (
            '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:rPr><w:noProof/></w:rPr>'
            '<w:pict>'
            '<v:shape xmlns:v="urn:schemas-microsoft-com:vml"'
            ' xmlns:o="urn:schemas-microsoft-com:office:office"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
            ' style="position:absolute;margin-left:0;margin-top:336pt;'
            'width:340pt;height:340pt;z-index:-251659264;'
            'mso-position-horizontal:center;mso-position-horizontal-relative:page;'
            'mso-position-vertical:absolute;mso-position-vertical-relative:page"'
            ' o:allowincell="f">'
            f'<v:imagedata r:id="{rId}" o:title="brasao"'
            ' gain="19661f" blacklevel="22938f"/>'
            '</v:shape>'
            '</w:pict>'
            '</w:r>'
        )
        wmk_p = OxmlElement("w:p")
        wmk_p.append(etree.fromstring(WMK_XML))
        hdr._element.insert(0, wmk_p)
    except Exception:
        pass


# ── Geração principal ──────────────────────────────────────────────────────────
def gerar(dados: dict) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin      = Cm(3.8)
        section.bottom_margin   = Cm(2.5)
        section.left_margin     = Cm(2)
        section.right_margin    = Cm(2)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)

    _add_header(doc)
    _add_footer(doc)
    _add_watermark(doc)

    # ── Título ──────────────────────────────────────────────────────────────
    titulo = doc.add_paragraph("CONTROLE DE RECEBIMENTO DE MATERIAIS")
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.paragraph_format.space_before = Pt(4)
    titulo.paragraph_format.space_after  = Pt(10)
    run_t = titulo.runs[0]
    run_t.bold = True
    run_t.font.size = Pt(14)
    run_t.font.color.rgb = _LARANJA
    run_t.font.underline = True

    # ── 1. Partes ────────────────────────────────────────────────────────────
    _info_table(doc, [
        ("CONTRATANTE", "Prefeitura Municipal de Maracanaú"),
        ("CONTRATADA",  dados.get("fornecedor", "")),
        ("CNPJ",        dados.get("cnpj", "")),
    ])
    _spacer(doc)

    # ── 2. Fundamentação legal ───────────────────────────────────────────────
    _info_table(doc, [
        ("PREGÃO DIGITAL",        dados.get("pregao",   "10.002/2024")),
        ("ATA DE REG. DE PREÇOS", dados.get("ata",      "10.004/2024")),
        ("CONTRATO",              dados.get("contrato", "")),
        ("EMPENHO",               dados.get("empenho",  "")),
    ], header="FUNDAMENTAÇÃO LEGAL")
    _spacer(doc)

    # ── 3. Objeto ────────────────────────────────────────────────────────────
    _info_table(doc, [
        ("", dados.get("objeto",
            "Aquisição de agregados, manufaturados e pré-moldados, para dar suporte "
            "às atividades da Secretaria de Infraestrutura, Mobilidade e "
            "Desenvolvimento Urbano do Município de Maracanaú-CE.")),
    ], header="OBJETO", col_widths=(Cm(0.1), Cm(16.9)))
    _spacer(doc)

    # ── 4. Dados do recebimento ───────────────────────────────────────────────
    _info_table(doc, [
        ("DANFE Nº",            dados.get("danfe", "")),
        ("DATA DE RECEBIMENTO", dados.get("data_recebimento", "")),
        ("LOCAL DE ENTREGA",    dados.get("local", "")),
    ], header="DADOS DO RECEBIMENTO")
    _spacer(doc)

    # ── 5. Tabela de itens ────────────────────────────────────────────────────
    items = dados.get("items", [])
    if items:
        sub = doc.add_paragraph("ITENS RECEBIDOS")
        sub.paragraph_format.space_before = Pt(2)
        sub.paragraph_format.space_after  = Pt(2)
        sub.runs[0].bold = True
        sub.runs[0].font.size = Pt(10)

        HDR    = ["Nº", "CÓDIGO", "DESCRIÇÃO", "MARCA", "UNIDADE", "QTDE", "VLR UNIT.", "VLR TOTAL"]
        WIDTHS = [Cm(0.8), Cm(1.6), Cm(6.5), Cm(2.0), Cm(1.7), Cm(1.4), Cm(1.5), Cm(1.5)]
        # soma = 17.0cm

        tbl = doc.add_table(rows=1 + len(items), cols=len(HDR))
        tbl.style = "Table Grid"
        tbl.autofit = False
        _lock_table_width(tbl, _TBL_W)
        for row in tbl.rows:
            for i, w in enumerate(WIDTHS):
                row.cells[i].width = w

        _header_row(tbl.rows[0], HDR)

        for idx, item in enumerate(items):
            row = tbl.rows[idx + 1]
            bg = "EBF3FB" if idx % 2 == 0 else "FFFFFF"
            for cell in row.cells:
                _set_bg(cell, bg)
            vals = [
                str(idx + 1),
                item.get("codigo", ""),
                item.get("descricao", ""),
                item.get("marca", ""),
                item.get("unidade", ""),
                item.get("quantidade", ""),
                item.get("valor_unitario", ""),
                item.get("valor_total", ""),
            ]
            for i, v in enumerate(vals):
                _cell(row.cells[i], v, center=(i != 2))

    _spacer(doc)

    # ── 6. Responsável pelo recebimento ───────────────────────────────────────
    resp  = dados.get("responsavel", "Clayton Rocha Braz")
    matr  = dados.get("matricula", "")
    cargo = dados.get("cargo", "DIRETOR PÁTIO DE MANUTENÇÃO – SEINFRA")

    tbl_resp = doc.add_table(rows=3, cols=2)
    tbl_resp.style = "Table Grid"
    tbl_resp.autofit = False
    _lock_table_width(tbl_resp, _TBL_W)
    tbl_resp.rows[0].cells[0].width = Cm(11.0)
    tbl_resp.rows[0].cells[1].width = Cm(6.0)

    # Cabeçalho da seção
    merged_hdr = tbl_resp.rows[0].cells[0].merge(tbl_resp.rows[0].cells[1])
    _set_bg(merged_hdr, _AZUL)
    _cell(merged_hdr, "RESPONSÁVEL PELO RECEBIMENTO", bold=True, white=True, center=True)

    # Linha NOME + MATRÍCULA
    _set_bg(tbl_resp.rows[1].cells[0], _CINZA_HDR)
    _set_bg(tbl_resp.rows[1].cells[1], _CINZA_HDR)
    _cell(tbl_resp.rows[1].cells[0], "NOME", bold=True)
    _cell(tbl_resp.rows[1].cells[1], "MATRÍCULA", bold=True)

    # Linha valores NOME + MATRÍCULA
    _cell(tbl_resp.rows[2].cells[0], resp)
    _cell(tbl_resp.rows[2].cells[1], matr)

    # Linha CARGO (mesclada)
    tbl_resp.add_row()
    tbl_resp.add_row()
    _set_bg(tbl_resp.rows[3].cells[0], _CINZA_HDR)
    merged_cargo_hdr = tbl_resp.rows[3].cells[0].merge(tbl_resp.rows[3].cells[1])
    _cell(merged_cargo_hdr, "CARGO", bold=True)

    merged_cargo_val = tbl_resp.rows[4].cells[0].merge(tbl_resp.rows[4].cells[1])
    _cell(merged_cargo_val, cargo)

    _spacer(doc, pts=8)

    # ── 7. Declaração ─────────────────────────────────────────────────────────
    decl = doc.add_paragraph(
        "Declaro, para os devidos fins, que os materiais e/ou serviços objeto deste "
        "documento foram recebidos nesta data, em conformidade com as especificações "
        "técnicas e quantitativas constantes no contrato/ordem de fornecimento, "
        "encontrando-se em perfeitas condições de uso e funcionamento."
    )
    decl.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    decl.paragraph_format.space_before = Pt(4)
    decl.paragraph_format.space_after  = Pt(20)
    decl.runs[0].italic = True
    decl.runs[0].font.size = Pt(8)

    # ── 8. Assinatura ─────────────────────────────────────────────────────────
    # Linha de assinatura
    p_linha = doc.add_paragraph()
    p_linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_linha.paragraph_format.space_before = Pt(0)
    p_linha.paragraph_format.space_after  = Pt(2)
    run_linha = p_linha.add_run("_" * 45)
    run_linha.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    for txt in (resp, cargo):
        p = doc.add_paragraph(txt)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
